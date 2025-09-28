#!/usr/bin/env python3
"""
Draft cover letters and tailor resumes using the upgraded resume flow.

- Uses src.tailor.resume.{call_llm_weaves, inject_skills, apply_complex_rewrites,
  apply_weaves_anywhere, fetch_jd_plaintext, canon}
  to create per-job tailored resumes (complex sentence rewrites + format-preserving weaving).
- Keeps the same output shape your UI expects:
  docs/<uid>/outbox/*.md
  docs/<uid>/resumes/*.docx
  docs/<uid>/changes/*.json  (object with company/title/paths/cover_meta/changes[])
  docs/<uid>/changes/*.jd.txt

Reads the shortlist from docs/data/scores.top.json (preferred, built by the GH step),
or falls back to docs/data/scores.json.

This script does NOT write drafts_index.json; your workflow step already builds it
by scanning docs/<uid>/* and uploads to Storage.
"""

import os, sys, json, re, yaml, hashlib, pathlib
from typing import Set, List, Dict, Optional

# repo root
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
sys.path.insert(0, ROOT)

from docx import Document

# covers
from src.tailor.cover import generate_cover_letter, get_company_context, pick_company_themes

# new resume flow helpers
from src.tailor.resume import (
    call_llm_weaves,
    inject_skills,
    apply_complex_rewrites,
    apply_weaves_anywhere,
    fetch_jd_plaintext,
    canon,
)

# vocab augmentation (existing)
from src.skills.taxonomy import augment_allowed_vocab

import requests
from bs4 import BeautifulSoup


# ------------------ misc helpers ------------------
def safe_slug(s: str) -> str:
    s = (s or "").strip()
    return "".join([c for c in s if c.isalnum() or c in ('-','_',' ')])[:150].strip().replace(" ", "_")

def norm(w: str) -> str:
    SYN = {
      "js":"javascript","reactjs":"react","ts":"typescript","ml":"machine learning",
      "cv":"computer vision","postgres":"postgresql","gh actions":"github actions",
      "gh-actions":"github actions","ci/cd":"ci","llm":"machine learning","rest":"rest api",
      "etl":"data pipeline"
    }
    w2 = (w or "").strip().lower()
    return SYN.get(w2, w2)

def tokens(text: str) -> Set[str]:
    WORD_RE = re.compile(r"[A-Za-z][A-Za-z0-9+./-]{1,}")
    return set(WORD_RE.findall((text or "").lower()))

STOPWORDS = {
    "engineer","engineering","software","developer","develop","team","teams","experience","years","year",
    "the","and","for","with","to","of","in","on","as","by","or","an","a","at","from","using",
    "we","you","our","your","will","work","role","responsibilities","requirements","preferred","must",
    "strong","plus","bonus","including","include","etc","ability","skills","excellent","communication",
}

def jd_sha(s: str) -> str:
    return hashlib.sha1((s or "").encode("utf-8")).hexdigest()[:8]


# ------------------ Supabase profile ------------------
SUPABASE_URL = os.environ.get("SUPABASE_URL","").rstrip("/")
SRK = os.environ.get("SUPABASE_SERVICE_ROLE_KEY","")

def load_profile_for_user(user_id: str) -> dict:
    if not (SUPABASE_URL and SRK and user_id):
        return {}
    url = f"{SUPABASE_URL}/rest/v1/profiles?id=eq.{user_id}&select=*"
    r = requests.get(url, headers={"apikey": SRK, "Authorization": f"Bearer {SRK}"}, timeout=30)
    r.raise_for_status()
    arr = r.json()
    return (arr[0] if arr else {}) or {}


# ------------------ allowed vocab ------------------
def _allowed_vocab_from_profile(profile: dict, portfolio: dict) -> Set[str]:
    skills = {str(s).lower() for s in (profile.get("skills") or [])}
    titles = {str(t).lower() for t in (profile.get("target_titles") or [])}
    tags = set()
    for section in ("projects", "work_experience", "workshops"):
        for item in (portfolio.get(section, []) or []):
            for b in (item.get("bullets", []) or []):
                for t in (b.get("tags", []) or []):
                    tags.add(str(t).lower())
    expanded = {norm(w) for w in (skills | tags | titles)}
    return set(expanded | skills | tags | titles)

def allowed_vocab(profile: dict, portfolio: dict) -> List[str]:
    base = _allowed_vocab_from_profile(profile, portfolio)
    titles = list(profile.get("target_titles") or [])
    return sorted(set(augment_allowed_vocab(base, titles)))


# ------------------ keyword scoring (for cover context / ATS list) ------------------
def _count_phrase(text: str, phrase: str) -> int:
    if not phrase:
        return 0
    pat = re.compile(rf"\b{re.escape(phrase)}\b", flags=re.IGNORECASE)
    return len(pat.findall(text or ""))

def extract_jd_terms(job: dict, allowed: set, cap=24) -> List[str]:
    title = (job.get("title") or "").lower()
    desc  = (job.get("description") or "").lower()
    url   = job.get("url", "")

    allowed_norm = {norm(a) for a in allowed}
    phrases = [a for a in allowed_norm if " " in a]
    unigrams = [a for a in allowed_norm if a and " " not in a]

    scores: Dict[str, float] = {}

    for ph in phrases:
        if any(x in STOPWORDS for x in ph.split()):
            continue
        c = _count_phrase(desc, ph)
        if c:
            scores[ph] = scores.get(ph, 0.0) + 3.0 * c
            if ph in title:
                scores[ph] += 2.0

    for w in list(tokens(desc)):
        w = norm(w)
        if w in unigrams and w not in STOPWORDS:
            scores[w] = scores.get(w, 0.0) + 1.0
            if w in title:
                scores[w] += 1.5

    lower_url = (url or "").lower()
    for k in list(scores.keys()):
        if k in lower_url:
            scores[k] += 0.5

    ranked = [k for k,_ in sorted(scores.items(), key=lambda kv: (-kv[1], kv[0]))]
    return ranked[:cap]


# ------------------ shortlist IO ------------------
def read_shortlist() -> List[dict]:
    """Prefer docs/data/scores.top.json (built by GH step); fallback to docs/data/scores.json."""
    top_path = os.path.join(ROOT, "docs", "data", "scores.top.json")
    base_path = os.path.join(ROOT, "docs", "data", "scores.json")
    if os.path.exists(top_path):
        with open(top_path) as f:
            data = json.load(f)
    elif os.path.exists(base_path):
        with open(base_path) as f:
            data = json.load(f)
    else:
        return []
    return data if isinstance(data, list) else \
           data.get("jobs") or data.get("items") or data.get("links") or []


def best_url(job: dict) -> str:
    return job.get("url") or job.get("link") or job.get("jd_url") or ""


def best_desc(job: dict) -> str:
    desc = (job.get("description") or "").strip()
    if len(desc) >= 800:
        return desc
    live = fetch_jd_plaintext(best_url(job)) if best_url(job) else ""
    return live if len(live) > len(desc) else desc


# ------------------ main ------------------
def main(top: int, user: Optional[str]):
    if not user:
        print("Missing --user; required for per-user output folders.")
        sys.exit(1)

    # Resolve per-user paths
    BASE_USER_DIR = os.path.join(ROOT, "docs", user)
    OUTBOX_MD   = os.path.join(BASE_USER_DIR, "outbox")
    RESUMES_MD  = os.path.join(BASE_USER_DIR, "resumes")
    CHANGES_DIR = os.path.join(BASE_USER_DIR, "changes")
    DATA_DIR    = os.path.join(ROOT, "docs", "data")
    PROFILE_YAML   = os.path.join(ROOT, "src", "core", "profile.yaml")
    PORTFOLIO_YAML = os.path.join(ROOT, "src", "core", "portfolio.yaml")

    # Load profile fresh (parse_resume step updated it)
    prof = load_profile_for_user(user)
    if prof:
        y = {
            "full_name": prof.get("full_name"),
            "email": prof.get("email"),
            "phone": prof.get("phone"),
            "skills": prof.get("skills") or [],
            "target_titles": prof.get("target_titles") or [],
            "locations": prof.get("locations") or [],
            "github": prof.get("github"),
        }
        os.makedirs(os.path.dirname(PROFILE_YAML), exist_ok=True)
        with open(PROFILE_YAML, "w") as f:
            yaml.safe_dump({k:v for k,v in y.items() if v is not None}, f)

    # Ensure portfolio file exists
    if not os.path.exists(PORTFOLIO_YAML):
        os.makedirs(os.path.dirname(PORTFOLIO_YAML), exist_ok=True)
        with open(PORTFOLIO_YAML, "w") as f:
            yaml.safe_dump({"projects": [], "work_experience": [], "workshops": []}, f)

    with open(PROFILE_YAML, 'r') as f:
        profile = yaml.safe_load(f) or {}
    with open(PORTFOLIO_YAML, 'r') as f:
        portfolio = yaml.safe_load(f) or {}

    allowed = set(allowed_vocab(profile, portfolio))

    # shortlist (Your Shortlist built earlier in the workflow)
    jobs: List[dict] = read_shortlist()
    if not jobs:
        print('No shortlist data found. Ensure the "Build TOP-N shortlist" step created docs/data/scores.top.json.')
        return
    # Keep order & dedupe by URL
    seen = set(); deduped = []
    for j in jobs:
        key = (best_url(j) or "").strip().lower() or f"no-url::{j.get('company','')}::{j.get('title','')}"
        if key in seen: continue
        seen.add(key); deduped.append(j)
    jobs = deduped[: max(1, min(20, int(top or 5)))]

    # dirs
    os.makedirs(OUTBOX_MD, exist_ok=True)
    os.makedirs(RESUMES_MD, exist_ok=True)
    os.makedirs(CHANGES_DIR, exist_ok=True)
    os.makedirs(DATA_DIR, exist_ok=True)

    BANLIST_JSON = os.path.join(DATA_DIR, 'banlist.json')
    try:
        with open(BANLIST_JSON, 'r') as bf:
            banlist = json.load(bf)
            if not isinstance(banlist, list): banlist = []
    except Exception:
        banlist = []
    banset = {x.strip().lower() for x in banlist}

    # choose base resume
    CURRENT_RESUME = os.path.join(ROOT, "assets", "current.docx")
    FALLBACK_RESUME= os.path.join(ROOT, "assets", "Resume-2025.docx")
    base_resume_path = CURRENT_RESUME if os.path.isfile(CURRENT_RESUME) else FALLBACK_RESUME
    print(f"Using base resume: {base_resume_path}")

    # summarize LLM usage
    use_llm = bool(os.getenv("OPENAI_API_KEY"))
    model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
    print(f"LLM for resume tailoring: {'enabled' if use_llm else 'disabled'} (model={model if use_llm else 'n/a'})")

    drafted_covers = drafted_resumes = 0
    llm_summary = {"used": use_llm, "model": (model if use_llm else None), "jobs": []}

    # We'll compute resume_plain once (from base resume) for planning
    base_doc_for_plain = Document(base_resume_path)
    resume_plain = "\n".join([p.text for p in base_doc_for_plain.paragraphs])

    # complex mode thresholds (mirrors src.tailor.resume envs so behavior is consistent)
    TAILOR_COMPLEX_MIN_BULLETS = int(os.getenv("TAILOR_COMPLEX_MIN_BULLETS", "6"))
    TAILOR_COMPLEX_MAX_BULLETS = int(os.getenv("TAILOR_COMPLEX_MAX_BULLETS", "10"))
    TAILOR_COMPLEX_MAX_WORDS   = int(os.getenv("TAILOR_COMPLEX_MAX_WORDS", "28"))
    TAILOR_MID_SENTENCE_STYLE  = os.getenv("TAILOR_MID_SENTENCE_STYLE", "comma")
    TAILOR_DASH_THRESHOLD      = int(os.getenv("TAILOR_DASH_THRESHOLD", "7"))
    TAILOR_END_PERIOD          = str(os.getenv("TAILOR_END_PERIOD", "1")).strip().lower() not in ("", "0", "false", "no")

    for j in jobs:
        company = j.get('company','') or j.get('org','')
        if company.lower() in banset:
            print(f"Skipping banned company: {company}")
            continue
        title   = j.get('title','') or j.get('job_title','')
        url     = best_url(j)

        safe_company = safe_slug(company)
        safe_title   = safe_slug(title)
        slug = f"{safe_company}_{safe_title}"[:150] or safe_slug(url) or "job"

        # ----- JD text -----
        jd_text = best_desc(j)
        tmp_job = dict(j); tmp_job["description"] = jd_text
        jd_kws = extract_jd_terms(tmp_job, allowed, cap=24)

        # save JD for UI (no hash in filename; see UI default)
        jd_txt_path = os.path.join(CHANGES_DIR, f"{slug}.jd.txt")
        with open(jd_txt_path, "w", encoding="utf-8") as f:
            f.write(jd_text[:20000])

        jd_hash = jd_sha(jd_text)

        # ----- cover letter (LLM-first; deterministic fallback handled in generate_cover_letter) -----
        ctx = get_company_context(j)
        company_themes = pick_company_themes(ctx)

        cover_fname = f"{slug}.md"
        cover_md = generate_cover_letter(
            job=j,
            profile=profile,
            jd_text=jd_text,
            jd_keywords=jd_kws,
            allowed_vocab=sorted(allowed),
            tone=os.getenv("COVER_TONE", "professional"),
        )
        with open(os.path.join(OUTBOX_MD, cover_fname), 'w', encoding="utf-8") as f:
            f.write(cover_md)
        j['cover_path'] = f"outbox/{cover_fname}"

        # ----- resume tailoring (complex rewrite + weave fallback) -----
        out_docx_name = f"{slug}_{jd_hash}.docx"
        out_docx = os.path.join(RESUMES_MD, out_docx_name)

        # fresh copy of the resume for this job
        doc = Document(base_resume_path)

        # core properties for ATS/debug
        try:
            cp = doc.core_properties
            cp.comments = f"job-copilot:{slug}:{jd_hash}"
            cp.subject = title
            # include up to 16 keywords for discoverability
            cp.keywords = ", ".join([k for k in jd_kws if k][:16])
        except Exception:
            pass

        # plan weaves with LLM (or deterministic fallback inside call_llm_weaves)
        plan = call_llm_weaves(resume_plain, jd_text, job_title=title, company=company)
        llm_phrases = [canon((w.get("phrase") or "").strip()) for w in (plan.get("weaves") or []) if (w.get("phrase") or "").strip()]

        # Apply skills (minimally, never creating new headers)
        skills_change = inject_skills(doc, plan.get("skills_additions") or [])
        granular_changes: List[Dict[str, str]] = []
        if skills_change:
            granular_changes.append({
                "anchor_section": skills_change["section"],
                "original_paragraph_text": skills_change["before"],
                "modified_paragraph_text": skills_change["after"],
                "inserted_sentence": None,
                "reason": skills_change["reason"],
            })

        # COMPLEX REWRITES (preferred path)
        complex_changes = apply_complex_rewrites(
            doc,
            jd_terms=(plan.get("jd_terms") or []) + llm_phrases,
            job_title=title,
            company=company,
            max_bullets=TAILOR_COMPLEX_MAX_BULLETS,
            style_hints={
                "mid_sentence_style": TAILOR_MID_SENTENCE_STYLE,
                "dash_threshold_words": TAILOR_DASH_THRESHOLD,
                "end_with_period": TAILOR_END_PERIOD,
                "max_words": TAILOR_COMPLEX_MAX_WORDS,
            },
        )
        granular_changes.extend(complex_changes)

        # LIGHT WEAVING (only if complex edits were few)
        if len(complex_changes) < TAILOR_COMPLEX_MIN_BULLETS:
            cands = (plan.get("skills_additions") or []) + llm_phrases
            default_phrase = canon(cands[0]) if cands else "requirements from the job description"
            weave_changes = apply_weaves_anywhere(doc, plan.get("weaves") or [], default_phrase)
            granular_changes.extend(weave_changes)

        # save tailored resume
        doc.save(out_docx)

        # ----- explain JSON with paths + cover meta (UI expects this object shape) -----
        explain = {
            "company": company,
            "title": title,
            "url": url,
            "ats_keywords": jd_kws,
            "llm_keywords": llm_phrases[:12],
            "changes": granular_changes,
            "jd_hash": jd_hash,
            "paths": {
                "resume_docx": f"resumes/{out_docx_name}",
                "cover_md": f"outbox/{cover_fname}",
                "jd_text": f"changes/{slug}.jd.txt",
            },
            "cover_meta": {
                "company_themes": company_themes[:8],
                "tone": os.getenv("COVER_TONE", "professional"),
            },
        }
        with open(os.path.join(CHANGES_DIR, f"{slug}_{jd_hash}.json"), 'w', encoding="utf-8") as f:
            json.dump(explain, f, ensure_ascii=False, indent=2)

        # summary stats
        llm_summary["jobs"].append({
            "slug": slug,
            "cover": True,
            "resume_injected": bool(granular_changes),
            "changes": len(granular_changes),
        })
        drafted_covers += 1
        drafted_resumes += 1

    # Persist a minimal mirror of the shortlist + LLM summary (informational)
    with open(os.path.join(ROOT, "docs", "data", "scores.json"), 'w', encoding="utf-8") as f:
        json.dump(jobs, f, ensure_ascii=False, indent=2)
    with open(os.path.join(ROOT, "docs", "data", "llm_info.json"), "w", encoding="utf-8") as f:
        json.dump(llm_summary, f, ensure_ascii=False, indent=2)
    with open(os.path.join(ROOT, "docs", "data", "banlist.json"), 'w', encoding="utf-8") as bf:
        json.dump(sorted(list(banset)), bf, ensure_ascii=False, indent=2)

    print(f"Drafted {drafted_covers} cover letters -> {OUTBOX_MD}")
    print(f"Drafted {drafted_resumes} tailored resumes -> {RESUMES_MD}")


if __name__ == '__main__':
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument('--top', type=int, default=5, help="Top-N from shortlist to draft (will still use docs/data/scores.top.json if present).")
    ap.add_argument('--user', type=str, required=True)
    args = ap.parse_args()
    main(args.top, args.user)
