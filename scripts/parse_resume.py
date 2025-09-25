# scripts/parse_resume.py
import os, re, json, hashlib, requests, argparse
from docx import Document
from typing import Iterable, Set

SUPABASE_URL = os.environ["SUPABASE_URL"].rstrip("/")
SRK = os.environ["SUPABASE_SERVICE_ROLE_KEY"]

PHONE_RE = re.compile(r"(?:\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}")
EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")

# Curated skills/terms. Keep tech here too (it’s fine), but we’ll match strictly.
KNOWN_TERMS = {
  # editorial / comms
  "editor", "editorial", "copyediting", "copy editing", "content editing", "proofreading",
  "ap style", "ap-style", "style guide", "publishing", "creative writing",
  # admin / office
  "microsoft office", "word", "excel", "powerpoint", "outlook", "adobe",
  "data entry", "record keeping", "scheduling",
  # general
  "customer service", "stakeholder management",
  # tech (kept, but matched strictly so false positives stop happening)
  "python","pytorch","opencv","javascript","typescript","react","react native","expo",
  "flask","sql","postgresql","plpgsql","linux","github actions","playwright",
  "hugging face","pwa","service worker","resnet","computer vision","c","c++","java","html","css"
}

WORD_RE = re.compile(r"[A-Za-z][A-Za-z0-9+.-]{1,}")  # tokens like "c++" ok

def tokenize(text: str) -> Set[str]:
  return {w.lower() for w in WORD_RE.findall(text or "")}

def has_all_words(skill: str, toks: Set[str]) -> bool:
  """Require exact whole-word presence for every word in the phrase."""
  parts = [p for p in re.split(r"\s+", skill.lower().strip()) if p]
  # treat c++ specially
  if skill.lower() == "c++":
    return "c++" in toks
  if len(parts) == 1:
    w = parts[0]
    # ignore pathological single-letter collisions (e.g. 'c' inside words)
    if len(w) == 1:
      return w in toks
    return w in toks
  return all(p in toks for p in parts)

def extract(docx_path: str):
  doc = Document(docx_path)
  full = "\n".join(p.text for p in doc.paragraphs)
  toks = tokenize(full)

  phone_m = PHONE_RE.search(full)
  email_m = EMAIL_RE.search(full)
  phone = phone_m.group(0) if phone_m else None
  email = email_m.group(0) if email_m else None

  skills = sorted({s for s in KNOWN_TERMS if has_all_words(s, toks)})

  # try to guess a name from the first few lines
  name = None
  for p in doc.paragraphs[:8]:
    t = (p.text or "").strip()
    if t and len(t.split()) <= 6 and not EMAIL_RE.search(t) and not PHONE_RE.search(t):
      name = t; break

  return {"full_name": name, "email": email, "phone": phone, "skills": skills}

def patch_profile(user_id: str, profile: dict):
  url = f"{SUPABASE_URL}/rest/v1/profiles?id=eq.{user_id}"
  r = requests.patch(
    url,
    headers={
      "Content-Type": "application/json",
      "apikey": SRK,
      "Authorization": f"Bearer {SRK}",
      "Prefer": "return=minimal",
    },
    data=json.dumps(profile),
    timeout=30,
  )
  r.raise_for_status()

def sha1_of(path: str) -> str:
  h = hashlib.sha1()
  with open(path, "rb") as f:
    for chunk in iter(lambda: f.read(8192), b""):
      h.update(chunk)
  return h.hexdigest()

def main(user_id: str):
  path = "assets/current.docx"  # NEW unified path
  if not os.path.exists(path):
    print("No resume at", path); return
  print("Parsing resume:", path, "sha1=", sha1_of(path))

  prof = extract(path)
  prof = {k: v for k, v in prof.items() if v is not None}

  if not prof:
    print("No fields extracted."); return

  patch_profile(user_id, prof)
  print("Patched profile (replaced fields: skills/name/email/phone).")

if __name__ == "__main__":
  ap = argparse.ArgumentParser(); ap.add_argument("--user", required=True)
  main(ap.parse_args().user)
