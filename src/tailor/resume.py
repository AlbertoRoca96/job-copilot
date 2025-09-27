#!/usr/bin/env python3
import os, sys, re, json, time, argparse, hashlib, pathlib, logging
from typing import List, Dict, Any, Optional, Tuple, Set
from copy import deepcopy

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ----------------------- config -----------------------
UA = "job-copilot/1.0 (+https://github.com/AlbertoRoca96/job-copilot)"
TIMEOUT = (10, 20)  # connect, read
MAX_JD_CHARS = 120_000
MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
TAILOR_INLINE_ONLY = str(os.getenv("TAILOR_INLINE_ONLY", "")).strip().lower() not in ("", "0", "false", "no")

# grammar/placement toggles
def _env_flag(name: str, default: bool = True) -> bool:
    return str(os.getenv(name, "1" if default else "0")).strip().lower() not in ("", "0", "false", "no")

def _env_opt(name: str, default: str) -> str:
    v = str(os.getenv(name, default)).strip().lower()
    return v or default

TAILOR_SMART_INSERT      = _env_flag("TAILOR_SMART_INSERT", True)
TAILOR_MID_SENTENCE_STYLE= _env_opt("TAILOR_MID_SENTENCE_STYLE", "comma")  # comma|dash|auto
TAILOR_DASH_THRESHOLD    = int(os.getenv("TAILOR_DASH_THRESHOLD", "7"))
TAILOR_CAP_SENTENCE      = _env_flag("TAILOR_CAP_SENTENCE", True)
TAILOR_END_PERIOD        = _env_flag("TAILOR_END_PERIOD", True)

logging.basicConfig(level=logging.INFO, format="%(levelname)s %(message)s")

# ----------------------- utils -----------------------
CANON = {
    "ap style": "AP style", "cms": "CMS", "pos": "POS", "crm": "CRM",
    "microsoft office": "Microsoft Office", "word": "Word", "excel": "Excel",
    "powerpoint": "PowerPoint", "outlook": "Outlook", "adobe": "Adobe",
    "photoshop": "Photoshop", "illustrator": "Illustrator", "indesign": "InDesign",
    "social media": "Social media", "content calendar": "Content calendar",
    "copyediting": "Copyediting", "fact checking": "Fact-checking",
    "proofreading": "Proofreading", "sql": "SQL", "supabase": "Supabase",
    "github actions": "GitHub Actions", "python": "Python", "javascript": "JavaScript"
}
WORD_RE = re.compile(r"[A-Za-z][A-Za-z0-9+.-]{1,}")

def normalize_ws(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip()

def canon(s: str) -> str:
    out = s or ""
    for k in sorted(CANON, key=len, reverse=True):
        out = re.sub(rf"\b{re.escape(k)}\b", CANON[k], out, flags=re.IGNORECASE)
    return out

def tokens(s: str) -> List[str]:
    return WORD_RE.findall((s or "").lower())

def token_set(s: str) -> Set[str]:
    return set(tokens(s))

def slugify(s: str) -> str:
    s = (s or "").lower()
    s = re.sub(r"[^a-z0-9]+", "-", s).strip("-")
    return s or "job"

# ----------------------- shortlist / links -----------------------
def read_links(path: str) -> List[Dict[str, Any]]:
    p = pathlib.Path(path)
    if not p.exists():
        raise FileNotFoundError(path)
    text = p.read_text(encoding="utf-8", errors="ignore")
    # try JSON first
    try:
        data = json.loads(text)
        out = []
        if isinstance(data, list):
            for item in data:
                if isinstance(item, str):
                    out.append({"url": item})
                elif isinstance(item, dict):
                    url = item.get("url") or item.get("link") or item.get("jd_url")
                    if url:
                        out.append({"url": url, **{k:v for k,v in item.items() if k not in ("url","link","jd_url")}})
        elif isinstance(data, dict):
            arr = data.get("jobs") or data.get("items") or data.get("links") or []
            for item in arr:
                if isinstance(item, str):
                    out.append({"url": item})
                elif isinstance(item, dict):
                    url = item.get("url") or item.get("link") or item.get("jd_url")
                    if url:
                        out.append({"url": url, **{k:v for k,v in item.items() if k not in ("url","link","jd_url")}})
        if out: return out
    except Exception:
        pass
    # else treat as newline-delimited list
    out = []
    for line in text.splitlines():
        line = line.strip()
        if line and re.match(r"^https?://", line):
            out.append({"url": line})
    return out

# ----------------------- JD fetch -----------------------
def fetch_jd_plaintext(url: str) -> str:
    """Fetch HTML and collapse to readable text (+ meta descriptions for gated sites)."""
    resp = requests.get(url, headers={
        "User-Agent": UA,
        "Accept": "text/html,application/xhtml+xml",
        "Accept-Language": "en-US,en;q=0.8",
    }, timeout=TIMEOUT, allow_redirects=True)
    resp.raise_for_status()
    html = resp.text
    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "noscript", "svg", "img"]):
        tag.decompose()
    for sel in ["header", "footer", "nav"]:
        for tag in soup.select(sel):
            tag.decompose()

    meta_bits = []
    for m in soup.find_all("meta", attrs={"name": "description"}):
        c = (m.get("content") or "").strip()
        if c: meta_bits.append(c)
    for m in soup.find_all("meta", attrs={"property": "og:description"}):
        c = (m.get("content") or "").strip()
        if c: meta_bits.append(c)

    text = soup.get_text(separator=" ", strip=True)
    blob = normalize_ws(" ".join([text] + meta_bits))
    return blob[:MAX_JD_CHARS]

# ----------------------- local keyword miner (fallback) -----------------------
PHRASES = [
    "ap style","cms","content calendar","social media","copyediting","fact checking",
    "proofreading","editorial calendar","pitching","seo","analytics","style guide",
    "calendar management","travel arrangements","expense reports","meeting notes",
    "inbox management","crm","data entry","pos","inventory","front of house",
    "customer service","microsoft office","excel","powerpoint","outlook","adobe",
    "photoshop","illustrator","indesign"
]

def mined_plan(resume_text: str, jd_text: str) -> Dict[str, Any]:
    jd_low = (jd_text or "").lower()
    res_low = (resume_text or "").lower()
    hits = [ph for ph in PHRASES if ph in jd_low]
    # dedupe keep-order
    seen, ordered = set(), []
    for h in hits:
        if h not in seen:
            ordered.append(h); seen.add(h)
    skills_additions = [canon(h) for h in ordered if h not in res_low][:6]

    weave_pool = []
    if "ap style" in jd_low or "cms" in jd_low:
        weave_pool.append("AP style and CMS guidelines")
    if "social media" in jd_low:
        weave_pool.append("social media scheduling and analytics")
    if "crm" in jd_low:
        weave_pool.append("CRM tracking and follow-ups")
    if "excel" in jd_low or "spreadsheet" in jd_low:
        weave_pool.append("Excel for tracking and reporting")
    if "pos" in jd_low or "inventory" in jd_low:
        weave_pool.append("POS and inventory checks")
    weaves = [{"section":"Work Experience","cue":"","phrase":p} for p in weave_pool[:4]]
    return {"skills_additions": skills_additions, "weaves": weaves}

# ----------------------- LLM (chat completions JSON mode) -----------------------
def call_llm_weaves(resume_text: str, jd_text: str, job_title: str = "", company: str = "") -> Dict[str, Any]:
    if not OPENAI_API_KEY:
        logging.warning("OPENAI_API_KEY not set; using deterministic fallback plan.")
        return mined_plan(resume_text, jd_text)
    try:
        from openai import OpenAI
        client = OpenAI(api_key=OPENAI_API_KEY)
        sys_prompt = (
            "You inject ATS-relevant keywords into an existing resume without fabricating achievements. "
            "Prefer weaving short prepositional phrases (e.g., 'with CRM tracking and follow-ups', "
            "'via POS and inventory checks') into bullets that already talk about the task. "
            "Keep phrases <= 18 words. Avoid buzzword stuffing. Output JSON only."
        )
        user_prompt = f"""Job title: {job_title or 'N/A'}
Company: {company or 'N/A'}

=== Job Description (plain text) ===
{jd_text}

=== Resume (plain text) ===
{resume_text}

Return JSON with exactly this shape:
{{
  "skills_additions": ["AP style", "CMS", "Excel"],
  "weaves": [
    {{"section":"Work Experience","cue":"edited","phrase":"AP style and CMS guidelines"}},
    {{"section":"Projects","cue":"wrote","phrase":"SEO keyword research and analytics"}}
  ]
}}"""
        resp = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "system", "content": sys_prompt},
                      {"role": "user", "content": user_prompt}],
            response_format={"type": "json_object"},
            temperature=0.2,
        )
        content = resp.choices[0].message.content or "{}"
        plan = json.loads(content)
    except Exception as e:
        logging.warning("LLM call failed (%s). Falling back to deterministic miner.", e)
        plan = {}

    if not isinstance(plan, dict): plan = {}
    skills_additions = [canon(x) for x in (plan.get("skills_additions") or []) if isinstance(x, str) and x.strip()]
    cleaned_weaves = []
    for w in (plan.get("weaves") or []):
        if not isinstance(w, dict): continue
        phrase = canon((w.get("phrase") or "").strip())
        if not phrase or len(phrase.split()) > 18: continue
        section = (w.get("section") or "Work Experience").strip()
        cue = (w.get("cue") or "").strip()
        cleaned_weaves.append({"section": section, "cue": cue, "phrase": phrase})
    if not skills_additions and not cleaned_weaves:
        return mined_plan(resume_text, jd_text)
    return {"skills_additions": skills_additions[:8], "weaves": cleaned_weaves[:6]}

# ----------------------- .docx helpers -----------------------
def paragraph_is_bullet(p: Paragraph) -> bool:
    # true Word bullets/numbering: presence of w:numPr on the paragraph
    try:
        pPr = p._p.pPr
        if (pPr is not None) and (pPr.numPr is not None):
            return True
    except Exception:
        pass
    # style hints:
    try:
        name = (getattr(p.style, "name", "") or "").lower()
    except Exception:
        name = ""
    if any(k in name for k in ("list","bullet","number")):
        return True
    # visible glyphs as last resort:
    t = normalize_ws(p.text)
    return t.startswith(("•","-","–","—","·"))

def dominant_run(p: Paragraph) -> Optional[Run]:
    best, best_len = None, -1
    for r in p.runs:
        txt = (r.text or "")
        style_name = (getattr(r.style, "name", "") or "").lower()
        if not txt.strip(): continue
        if "hyperlink" in style_name: continue
        L = len(txt.strip())
        if L > best_len:
            best, best_len = r, L
    if best is None:
        for r in reversed(p.runs):
            if (r.text or "").strip():
                return r
    return best

def copy_format(src: Run, dst: Run):
    try:
        dst.font.name = src.font.name
        dst.font.size = src.font.size
        dst.font.bold = src.font.bold
        dst.font.italic = src.font.italic
        dst.font.underline = src.font.underline
        try: dst.style = src.style
        except Exception: pass
    except Exception:
        pass

# ---- XML-level helpers (for exact formatting preservation) ----
def el_text(p_el) -> str:
    return "".join(t.text for t in p_el.xpath('.//*[local-name()="t"]') if t.text)

def el_is_bullet(p_el) -> bool:
    return bool(p_el.xpath('./*[local-name()="pPr"]/*[local-name()="numPr"]'))

def el_clone_rpr_from(src_r) -> Optional[OxmlElement]:
    try:
        rPr = getattr(src_r, "rPr", None)
        return deepcopy(rPr) if rPr is not None else None
    except Exception:
        return None

def el_last_meaningful_r(p_el):
    rs = p_el.xpath('./*[local-name()="r"]')
    for r in reversed(rs):
        texts = r.xpath('.//*[local-name()="t"]')
        if any((t.text or "").strip() for t in texts):
            return r
    return rs[-1] if rs else None

def el_append_run(p_el, text: str):
    base = el_last_meaningful_r(p_el)
    r = OxmlElement('w:r')
    rPr = el_clone_rpr_from(base) if base is not None else None
    if rPr is not None:
        r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p_el.append(r)

def xml_all_paragraphs(doc: Document):
    return list(doc.element.xpath('.//*[local-name()="p"]'))

# ----------------------- grammar bridge -----------------------
_LEADS_TO_KEEP = ("to ", "for ")
_STRIP_LEADS = (
    "using ", "with ", "via ", "through ", "by ",
    "while ", "as ", "as part of ", "in order to ", "in accordance with ",
    "per "
)
_GERUNDS = ("managing", "coordinating", "handling", "scheduling", "maintaining",
            "performing", "providing", "ensuring", "tracking", "verifying",
            "triaging", "documenting", "updating", "supporting", "resolving")

def _sentence_case(s: str) -> str:
    for i, ch in enumerate(s):
        if ch.isalpha():
            return s[:i] + ch.upper() + s[i+1:]
    return s

def _choose_mid_delim(bridge: str) -> str:
    style = TAILOR_MID_SENTENCE_STYLE
    if style == "comma":
        return ", "
    if style == "dash":
        return " — "
    # auto
    n = len((bridge or "").split())
    return " — " if n >= TAILOR_DASH_THRESHOLD else ", "

def bridge_phrase(raw: str) -> str:
    """
    Normalize a phrase and choose a natural connector.
    Returns text WITHOUT leading space; caller adds spacing.
    """
    p = canon((raw or "").strip())
    p = p.rstrip(". ")
    low = p.lower()

    # strip awkward leads
    for lead in _STRIP_LEADS:
        if low.startswith(lead):
            p = p[len(lead):]
            low = p.lower()
            break

    # keep if already "to ..." or "for ..."
    for keep in _LEADS_TO_KEEP:
        if low.startswith(keep):
            return p

    if any(low.startswith(g+" ") for g in _GERUNDS):
        return f"by {p}"
    # default nouny case
    return f"with {p}"

# ----------------------- section detection -----------------------
SECTION_NAMES = [
    "work experience", "experience", "professional experience",
    "employment", "relevant experience"
]

def find_section_ranges(doc: Document, titles: List[str]) -> Dict[str, Tuple[int,int]]:
    wants = [normalize_ws(t).lower() for t in titles]
    hits: Dict[str,int] = {}
    for i,p in enumerate(doc.paragraphs):
        if normalize_ws(p.text).lower() in wants:
            hits[normalize_ws(p.text).lower()] = i
    ranges: Dict[str,Tuple[int,int]] = {}
    for k,start in sorted(hits.items(), key=lambda kv: kv[1]):
        later = [v for _,v in sorted(hits.items(), key=lambda kv: kv[1]) if v>start]
        end = later[0] if later else len(doc.paragraphs)
        ranges[k]=(start,end)
    return ranges

def work_experience_start(doc: Document) -> int:
    ranges = find_section_ranges(doc, SECTION_NAMES)
    for name in SECTION_NAMES:
        k = name
        if k in ranges: return ranges[k][0]
    # fallback: skip top 8 paragraphs (headers/summary region)
    return min(len(doc.paragraphs), 8)

# ----------------------- candidate discovery -----------------------
def candidate_paragraphs(doc: Document) -> List[Paragraph]:
    start_idx = work_experience_start(doc)
    bullets = []
    bodies = []
    for i, p in enumerate(doc.paragraphs):
        if i < start_idx:  # keep above-work-experience pristine
            continue
        t = normalize_ws(p.text)
        if not t:
            continue
        if paragraph_is_bullet(p):
            bullets.append((i, p))
        elif len(t) >= 25:
            bodies.append((i, p))
    # prefer bullets, then longer bodies
    return [p for _, p in bullets] + [p for _, p in bodies]

def candidate_paragraph_els(doc: Document):
    p_els = xml_all_paragraphs(doc)
    # approximate start index by matching the Nth python-docx paragraph element
    start = work_experience_start(doc)
    # map python paragraphs to xml nodes by text match (best-effort)
    # simpler: just skip first ~start*2 xml p's as a heuristic
    return p_els[max(0, start):]

# ----------------------- low-level run insertion (format-safe) -----------------------
def _insert_run_after(p: Paragraph, after_run: Run, text: str, style_src: Optional[Run] = None) -> Run:
    """
    Insert a run directly after `after_run` (preserving formatting).
    """
    base_el = after_run._r
    new_r = OxmlElement('w:r')
    src = style_src or after_run
    try:
        rPr = deepcopy(src._r.rPr) if getattr(src._r, "rPr", None) is not None else None
    except Exception:
        rPr = None
    if rPr is not None:
        new_r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    new_r.append(t)
    base_el.addnext(new_r)
    return Run(new_r, p)

# ----------------------- injectors -----------------------
_TRAILING_RE = re.compile(
    r"\b(including|using|via|through|resulting in|while|which|that)\b", re.IGNORECASE
)

def _find_trailing_spot(text: str) -> Optional[int]:
    """
    Find a reasonable point to insert BEFORE trailing add-ons.
    Returns absolute char index or None.
    """
    if not TAILOR_SMART_INSERT:
        return None
    m = _TRAILING_RE.search(text or "")
    return m.start() if m else None

def insert_run_at_end(p: Paragraph, bridge: str) -> Tuple[bool, str, str, str]:
    """
    Grammar-aware append that preserves formatting.
    Returns (ok, before, after, inserted_core_text).
    """
    base = dominant_run(p)
    before = "".join(r.text for r in p.runs) if p.runs else p.text
    if not before.strip():
        return (False, before, before, "")
    if bridge.lower() in before.lower():
        return (False, before, before, "")

    # Try smart mid-sentence insertion BEFORE trailing phrases (format-safe split)
    cut = _find_trailing_spot(before)
    if cut is not None and p.runs:
        # choose delimiter
        delim = _choose_mid_delim(bridge)
        insertion = f"{delim}{bridge}"

        # locate run containing `cut`
        acc = 0
        for r in p.runs:
            txt = r.text or ""
            nxt = acc + len(txt)
            if cut <= nxt:
                local = cut - acc
                try:
                    left, right = txt[:local], txt[local:]
                    r.text = left
                    ins_run = _insert_run_after(p, r, insertion, style_src=base or r)
                    _insert_run_after(p, ins_run, right, style_src=r)
                    after = "".join(rr.text for rr in p.runs)
                    return (True, before, after, insertion.strip())
                except Exception:
                    # fall back to end-append
                    break
            acc = nxt

    # End-of-paragraph path (no risky mid-sentence surgery)
    trimmed = before.rstrip()
    end_ch = trimmed[-1] if trimmed else ""
    if end_ch in ".!?":
        # New sentence: Capitalize + (optional) trailing period.
        insertion_core = bridge
        if TAILOR_CAP_SENTENCE:
            insertion_core = _sentence_case(insertion_core)
        inserted_text = (" " + insertion_core + ( "." if TAILOR_END_PERIOD and not insertion_core.endswith(".") else "" ))
    else:
        # Continue current sentence with comma or dash.
        delim = _choose_mid_delim(bridge)
        inserted_text = f"{delim}{bridge}"

    r = p.add_run(inserted_text)
    if base:
        copy_format(base, r)
    after = before + inserted_text
    return (True, before, after, inserted_text.strip())

def weave_into_el(p_el, bridge: str) -> Tuple[bool, str, str, str]:
    before = el_text(p_el)
    if not before.strip():
        return (False, before, before, "")
    if bridge.lower() in before.lower():
        return (False, before, before, "")

    # XML path: stick to safe end-append; grammar-aware spacing/punct only.
    trimmed = before.rstrip()
    end_ch = trimmed[-1] if trimmed else ""
    if end_ch in ".!?":
        insertion_core = bridge
        if TAILOR_CAP_SENTENCE:
            # naive sentence case for XML path
            insertion_core = _sentence_case(insertion_core)
        prefix = " "
        suffix = "." if TAILOR_END_PERIOD and not insertion_core.endswith(".") else ""
        ins_text = f"{prefix}{insertion_core}{suffix}"
    else:
        ins_text = f"{_choose_mid_delim(bridge)}{bridge}"

    el_append_run(p_el, ins_text)
    after = el_text(p_el)
    return (True, before, after, ins_text.strip())

# cue-scoring
def cue_score(text: str, cue: str) -> int:
    if not cue: return 0
    toks = token_set(text)
    cue_toks = token_set(cue)
    return sum(1 for c in cue_toks if c in toks)

# ----------------------- weaving helpers -----------------------
def apply_weaves_xml(doc: Document, weaves: List[Dict[str,str]], default_phrase: str) -> List[Dict[str,str]]:
    changes: List[Dict[str,str]] = []
    p_els = candidate_paragraph_els(doc)
    bullets = [(i, p_el) for i, p_el in enumerate(p_els) if el_is_bullet(p_el) and normalize_ws(el_text(p_el))] \
              or [(i, p_el) for i, p_el in enumerate(p_els) if len(normalize_ws(el_text(p_el))) >= 25]

    used = set()
    phrases = [w.get("phrase") for w in (weaves or []) if w.get("phrase")] or []
    cues    = [w.get("cue","") for w in (weaves or [])] or []
    if default_phrase and default_phrase not in phrases:
        phrases.append(default_phrase); cues.append("")

    inserted_any = False
    for idx, phrase in enumerate(phrases[:4]):
        bridge = bridge_phrase(phrase)
        # pick best bullet by cue score
        best = None
        best_score = -1
        for k, (_, p_el) in enumerate(bullets):
            if k in used: continue
            t = normalize_ws(el_text(p_el))
            s = cue_score(t, cues[idx] if idx < len(cues) else "")
            if s > best_score:
                best_score = s
                best = (k, p_el, t)
        if best:
            k, p_el, _ = best
            ok, before, after, ins = weave_into_el(p_el, bridge)
            if ok:
                used.add(k); inserted_any = True
                changes.append({
                    "anchor_section": "Work Experience",
                    "original_paragraph_text": before,
                    "modified_paragraph_text": after,
                    "inserted_sentence": ins,
                    "reason": "Injected inline JD keyword phrase (XML; formatting preserved)."
                })

    if not inserted_any and default_phrase:
        for p_el in p_els:
            before = normalize_ws(el_text(p_el))
            if before:
                ok, b2, a2, ins = weave_into_el(p_el, bridge_phrase(default_phrase))
                if ok:
                    changes.append({
                        "anchor_section": "Work Experience",
                        "original_paragraph_text": b2,
                        "modified_paragraph_text": a2,
                        "inserted_sentence": ins,
                        "reason": "Hard fallback to ensure visible tailoring (XML; formatting preserved)."
                    })
                break
    return changes

def apply_weaves_inline(doc: Document, weaves: List[Dict[str,str]], default_phrase: str) -> List[Dict[str,str]]:
    """
    Paragraph-level weaving WITHOUT rewriting the paragraph:
    we append a styled run so original formatting is untouched.
    """
    changes: List[Dict[str,str]] = []
    cands = candidate_paragraphs(doc)
    phrases = [w.get("phrase") for w in (weaves or []) if w.get("phrase")] or []
    cues    = [w.get("cue","") for w in (weaves or [])] or []
    if default_phrase and default_phrase not in phrases:
        phrases.append(default_phrase); cues.append("")

    used_idx: Set[int] = set()
    for idx, phrase in enumerate(phrases[:4]):
        bridge = bridge_phrase(phrase)
        # choose best candidate by cue score
        best_i, best_p, best_score = -1, None, -1
        for i, p in enumerate(cands):
            if i in used_idx: continue
            t = normalize_ws(p.text)
            if not t: continue
            s = cue_score(t, cues[idx] if idx < len(cues) else "")
            if s > best_score:
                best_i, best_p, best_score = i, p, s
        if best_p is None:
            continue
        ok, before, after, ins = insert_run_at_end(best_p, bridge)
        if ok:
            used_idx.add(best_i)
            changes.append({
                "anchor_section": "Work Experience",
                "original_paragraph_text": before,
                "modified_paragraph_text": after,
                "inserted_sentence": ins,
                "reason": "Injected inline JD keyword phrase (paragraph; formatting preserved)."
            })

    if not changes and default_phrase and cands:
        p = cands[0]
        ok, before, after, ins = insert_run_at_end(p, bridge_phrase(default_phrase))
        if ok:
            changes.append({
                "anchor_section": "Work Experience",
                "original_paragraph_text": before,
                "modified_paragraph_text": after,
                "inserted_sentence": ins,
                "reason": "Hard fallback to ensure visible tailoring (paragraph; formatting preserved)."
            })
    return changes

# ----------------------- skills: find + inject (unchanged API) -----------------------
def inject_skills(doc: Document, additions: List[str]) -> Optional[Dict[str,str]]:
    if not additions:
        return None
    additions = [canon(a) for a in additions if a]
    ranges = find_section_ranges(doc, ["Technical Skills", "Skills", "Core Skills"])
    if not ranges:
        return None
    key = next(iter([k for k in ("technical skills","skills","core skills") if k in ranges]), None)
    if not key:
        return None
    s, e = ranges[key]
    for i in range(s, e):
        p = doc.paragraphs[i]
        t = normalize_ws(p.text)
        if not t or paragraph_is_bullet(p):
            continue
        if "," not in t and ";" not in t:
            continue
        present = token_set(t)
        new = [a for a in additions if a.lower() not in present]
        if not new:
            return None
        before = "".join(r.text for r in p.runs) if p.runs else p.text
        t2 = t[:-1] if t.endswith(".") else t
        if ("," in t2 or ";" in t2):
            sep = ", " if not t2.endswith(",") else ""
            appended = f"{sep}{', '.join(new)}"
            after = t2 + appended
        else:
            after = t2 + f": {', '.join(new)}"
        # append as styled run to preserve inline formatting of existing text
        base = dominant_run(p)
        if p.text != t2:
            pass
        run = p.add_run(after[len(t2):])
        if base:
            copy_format(base, run)
        return {
            "section": "Skills/Technical Skills",
            "before": before,
            "after": "".join(r.text for r in p.runs),
            "inserted": None,
            "reason": "Enriched inline skills list (formatting preserved)."
        }
    return None

# ----------------------- orchestrator -----------------------
def apply_weaves_anywhere(doc: Document, weaves: List[Dict[str,str]], default_phrase: str) -> List[Dict[str,str]]:
    if TAILOR_INLINE_ONLY:
        logging.info("TAILOR_INLINE_ONLY=1 → using paragraph-level weaving (styled run append).")
        changes = apply_weaves_inline(doc, weaves, default_phrase)
        if not changes:
            logging.info("Paragraph weaving produced no changes; nothing to record.")
        return changes

    logging.info("Using XML-level weaving (with paragraph fallback).")
    changes = apply_weaves_xml(doc, weaves, default_phrase)
    if not changes:
        logging.info("XML weaving produced no changes; falling back to paragraph weaving.")
        changes = apply_weaves_inline(doc, weaves, default_phrase)
    return changes

# ----------------------- pipeline -----------------------
def run_pipeline(links_file: str, resume_path: str, out_prefix: str, uid: str = "user") -> None:
    out_root = pathlib.Path(out_prefix).joinpath(uid)
    out_resumes = out_root / "resumes"
    out_changes = out_root / "changes"
    out_outbox = out_root / "outbox"
    for d in (out_resumes, out_changes, out_outbox):
        d.mkdir(parents=True, exist_ok=True)

    resume_doc = Document(resume_path)
    resume_plain = "\n".join([p.text for p in resume_doc.paragraphs])

    links = read_links(links_file)
    if not links:
        raise RuntimeError("No links found.")

    index_items = []
    for item in links:
        url = item.get("url")
        if not url:
            continue
        logging.info("Fetching JD: %s", url)
        try:
            jd_text = fetch_jd_plaintext(url)
        except Exception as e:
            logging.warning("Failed to fetch %s: %s", url, e)
            continue

        job_title = item.get("title") or item.get("job_title") or ""
        company   = item.get("company") or item.get("org") or ""
        plan = call_llm_weaves(resume_plain, jd_text, job_title, company)

        # Save the plan for debugging (not used by the UI)
        slug_base = slugify(job_title or url)[:80]
        (out_changes / f"{slug_base}_plan.json").write_text(
            json.dumps(plan, ensure_ascii=False, indent=2), encoding="utf-8")
        # Apply to a fresh copy of the resume per job
        doc = Document(resume_path)
        changes: List[Dict[str, Any]] = []

        # Pass A: skills additions (kept minimal; no new headers ever)
        skills_change = inject_skills(doc, plan.get("skills_additions") or [])
        if skills_change:
            changes.append({
                "anchor_section": skills_change["section"],
                "original_paragraph_text": skills_change["before"],
                "modified_paragraph_text": skills_change["after"],
                "inserted_sentence": None,
                "reason": skills_change["reason"]
            })

        # Pass B: weaving (experience and below)
        cands = (plan.get("skills_additions") or []) + [w["phrase"] for w in (plan.get("weaves") or []) if w.get("phrase")]
        default_phrase = canon(cands[0]) if cands else "requirements from the job description"
        weave_changes = apply_weaves_anywhere(doc, plan.get("weaves") or [], default_phrase)
        changes.extend(weave_changes)

        # Persist artifacts
        jd_txt_path = out_changes / f"{slug_base}.jd.txt"
        jd_txt_path.write_text(jd_text, encoding="utf-8")

        h = hashlib.sha1((url + str(time.time())).encode()).hexdigest()[:8]
        out_docx = out_resumes / f"{slug_base}_{h}.docx"
        doc.save(out_docx.as_posix())

        changes_path = out_changes / f"{slug_base}_{h}.json"
        changes_path.write_text(json.dumps(changes, ensure_ascii=False, indent=2), encoding="utf-8")

        index_items.append({
            "url": url,
            "title": job_title,
            "company": company,
            "resume_path": str(out_docx),
            "jd_text_path": str(jd_txt_path),
            "changes_path": str(changes_path),
            "ts": int(time.time())
        })

    # Write drafts index for UI (exclude *_plan.json here)
    index_path = out_root / "drafts_index.json"
    index_path.write_text(json.dumps({
        "outbox": [""],
        "resumes": [pathlib.Path(i["resume_path"]).name for i in index_items],
        "changes": [pathlib.Path(i["changes_path"]).name for i in index_items],
    }, ensure_ascii=False, indent=2), encoding="utf-8")
    logging.info("Wrote index: %s", index_path)

# ----------------------- CLI -----------------------
def main():
    ap = argparse.ArgumentParser(description="Tailor resume from JD links with LLM-assisted weaving (format-preserving, Experience+ only).")
    ap.add_argument("--links", required=True, help="Path to JSON or .txt containing JD links.")
    ap.add_argument("--resume", required=True, help="Path to source .docx resume.")
    ap.add_argument("--out", required=True, help="Output prefix (e.g., outputs).")
    ap.add_argument("--user", default="user", help="User id folder under output prefix.")
    args = ap.parse_args()
    run_pipeline(args.links, args.resume, args.out, uid=args.user)

if __name__ == "__main__":
    main()
